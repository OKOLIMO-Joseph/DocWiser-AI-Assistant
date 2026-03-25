import PyPDF2
import docx
import io
from typing import Optional

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num} ---\n"
                    text += page_text + "\n"
            
            if not text.strip():
                return "No text could be extracted from PDF. The PDF might be scanned or image-based."
            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += f"\n{cell_text}"
            
            if not text.strip():
                return "No text could be extracted from Word document. The document might be empty."
            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: bytes, file_extension: str) -> str:
        """Extract text based on file type"""
        if file_extension == '.pdf':
            return DocumentParser.extract_text_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            return DocumentParser.extract_text_from_docx(file_content)
        else:
            raise Exception(f"Unsupported file type: {file_extension}")